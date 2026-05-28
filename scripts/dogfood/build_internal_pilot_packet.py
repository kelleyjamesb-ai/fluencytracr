#!/usr/bin/env python3
"""Build the internal FluencyTracr pilot packet from aggregate dogfood outputs.

This script reads retained aggregate CSV/readout artifacts only. It does not
read raw event rows, prompts, outputs, transcripts, or person-level exports.
"""

from __future__ import annotations

import argparse
import csv
import json
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]

GLOSSARY = [
    (
        "AI Work Evidence",
        "Aggregate telemetry that shows where AI is present in real work, how it is connected to workflows, and where evidence is strong or weak.",
    ),
    (
        "AI surface",
        "A place where AI can be used, such as search, chat, autocomplete, a workflow run, an agent span, or an embedded assistant.",
    ),
    (
        "Workflow evidence",
        "Aggregate signs that AI is attached to a real work path instead of a standalone interaction.",
    ),
    (
        "Velocity",
        "How quickly AI-assisted work appears to move, using aggregate frequency, engagement, and repeat-use patterns.",
    ),
    (
        "Depth",
        "How embedded AI appears to be across surfaces and workflows, not how skilled any individual person is.",
    ),
    (
        "Frequency",
        "How often AI activity appears in the aggregate window.",
    ),
    (
        "Engagement",
        "How much aggregate interaction occurs around AI-assisted work, such as repeated use or continued activity.",
    ),
    (
        "Breadth",
        "How many distinct AI surfaces or work contexts appear in the aggregate evidence.",
    ),
    (
        "Reliability Factor",
        "A confidence lens that looks for abandonment, friction loops, recovery, and verification signals before interpreting behavior.",
    ),
    (
        "Quality Multiplier",
        "A caveat lens that can discount or amplify time-saved assumptions later; this packet does not use it to calculate ROI.",
    ),
    (
        "Trust evidence",
        "Aggregate signs that people verified, continued, corrected, recovered, or gave feedback after AI assistance.",
    ),
    (
        "Trust attribution",
        "Whether a trust signal can be attached to a workflow path with enough confidence to interpret it.",
    ),
    (
        "Evidence gap",
        "A place where activity exists but the supporting source coverage, attribution, or downstream context is not strong enough to interpret.",
    ),
    (
        "Motif tier",
        "A plain-language bucket that separates raw activity from stronger workflow, verification, or source-coverage evidence.",
    ),
    (
        "Value-readiness zone",
        "An executive action lane that says what to do next, such as scale, redesign, repair trust loops, or hold interpretation.",
    ),
    (
        "Outcome evidence",
        "Customer-owned aggregate business context, such as a KPI window, that can be attached later without claiming causality by default.",
    ),
    (
        "Value hypothesis",
        "A testable idea about where AI may create value, such as acceleration, quality premium, or net-new work. It is not ROI proof.",
    ),
    (
        "Source coverage",
        "How complete and unambiguous the available telemetry is for a given aggregate interpretation.",
    ),
    (
        "Suppressed",
        "A fail-closed result where FluencyTracr withholds interpretation because the governed evidence bar is not met.",
    ),
]


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="") as handle:
        return list(csv.DictReader(handle))


def num(value: str | None, default: float = 0.0) -> float:
    if value is None or value == "":
        return default
    try:
        return float(value)
    except ValueError:
        return default


def fmt_int(value: float | int) -> str:
    return f"{int(round(value)):,}"


def fmt_pct(value: float) -> str:
    return f"{value * 100:.1f}%"


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def load_summary(input_dir: Path) -> dict:
    fresh = read_csv(input_dir / "episode_journey_microcosm_full_run.aggregate.csv")
    strategy = read_csv(ROOT / "dogfood-output/v4-value-realization-strategy/v4_value_realization_strategy_summary.csv")
    trust = read_csv(ROOT / "dogfood-output/v4-trust-attribution-refinement/v4_trust_attribution_summary.csv")
    trust_class = read_csv(ROOT / "dogfood-output/v4-trust-cohort-classification/v4_trust_classification_summary_safe.csv")
    skill = read_csv(ROOT / "dogfood-output/v4-skill-read-availability/skill_read_availability_all_windows.csv")
    outcome = read_csv(ROOT / "dogfood-output/v4-outcome-join-test/v4_outcome_join_readiness.csv")
    tsdr = read_csv(ROOT / "dogfood-output/v4-time-saved-defensibility-test/v4_time_saved_defensibility_readiness.csv")

    framework = []
    for row in fresh:
        if row["output_section"] == "FRAMEWORK_MICROCOSM_SUMMARY":
            framework.append(
                {
                    "window": row["window_id"].replace("_", " ").title(),
                    "sampled_windows": int(num(row["aggregate_count"])),
                    "avg_frequency": num(row["metric_1"]),
                    "avg_engagement": num(row["metric_2"]),
                    "avg_breadth": num(row["metric_3"]),
                    "avg_reliability": num(row["metric_4"]),
                    "avg_quality": num(row["metric_5"]),
                }
            )

    motif_rows = [row for row in fresh if row["output_section"] == "MOTIF_TIER_SUMMARY"]
    motif_by_window: dict[str, list[dict]] = defaultdict(list)
    motif_totals: dict[str, float] = defaultdict(float)
    for row in motif_rows:
        count = num(row["aggregate_count"])
        item = {
            "window": row["window_id"].replace("_", " ").title(),
            "tier": row["primary_dimension"],
            "count": count,
            "motif_shapes": num(row["metric_1"]),
            "avg_linked_events": num(row["metric_2"]),
            "friction_rate": num(row["metric_3"]),
            "continuation_rate": num(row["metric_4"]),
            "verification_rate": num(row["metric_5"]),
        }
        motif_by_window[item["window"]].append(item)
        motif_totals[item["tier"]] += count

    strategy_posture: dict[str, dict[str, float]] = defaultdict(lambda: {"rows": 0, "cohort": 0, "signals": 0, "attributed": 0})
    zone_summary: dict[str, dict[str, float]] = defaultdict(lambda: {"rows": 0, "cohort": 0, "signals": 0, "attributed": 0})
    for row in strategy:
        posture = row.get("strategy_posture", "UNKNOWN")
        zone = row.get("readout_zone_test_result", "UNKNOWN")
        row_count = num(row.get("row_count"))
        cohort = num(row.get("total_cohort_rows"))
        signals = num(row.get("total_signal_count"))
        attributed = num(row.get("total_attributed_signal_count"))
        for bucket, key in [(strategy_posture, posture), (zone_summary, zone)]:
            bucket[key]["rows"] += row_count
            bucket[key]["cohort"] += cohort
            bucket[key]["signals"] += signals
            bucket[key]["attributed"] += attributed

    trust_summary: dict[str, dict[str, float]] = defaultdict(lambda: {"signals": 0, "attributed": 0})
    for row in trust:
        key = row.get("attribution_status", "UNKNOWN")
        trust_summary[key]["signals"] += num(row.get("signal_count"))
        trust_summary[key]["attributed"] += num(row.get("attributed_signal_count"))

    trust_readout: dict[str, dict[str, float]] = defaultdict(lambda: {"signals": 0, "attributed": 0})
    for row in trust_class:
        key = row.get("readout_decision", "UNKNOWN")
        trust_readout[key]["signals"] += num(row.get("signal_count"))
        trust_readout[key]["attributed"] += num(row.get("attributed_signal_count"))

    skill_rows = []
    skill_total = 0.0
    skill_with_parent = 0.0
    for row in skill:
        count = num(row.get("skill_read_rows"))
        skill_total += count
        if row.get("parent_join_status") == "has_parent_join_key":
            skill_with_parent += count
        skill_rows.append(row)

    outcome_checks = {row["readiness_check"]: row for row in outcome}
    tsdr_checks = {row["readiness_check"]: row for row in tsdr}

    trust_gap = {
        "public_gap_episodes": 37_959_260,
        "public_gap_share": 0.431,
        "true_downstream_gap": 37_484_844,
        "ambiguous_boundary_fold_in": 474_414,
        "high_confidence_episode_share": 0.9995,
        "recovered_after_failure_share": 0.1799,
        "high_confidence_episodes": 87_985_613,
        "citation_available_high_confidence": 12_386_039,
        "citation_click_high_confidence": 509_226,
        "explicit_feedback_high_confidence": 64_681,
    }

    return {
        "generated_on": date.today().isoformat(),
        "fresh_job": {
            "job_id": "scio-apps:bqjob_r68a7577f065e48e5_0000019e6ee2c806_1",
            "duration": "0:06:00.970000",
            "bytes_processed": 4_926_181_266_149,
            "source": "scrubbed GCE customer events + scrubbed agent spans",
        },
        "framework": framework,
        "motif_by_window": motif_by_window,
        "motif_totals": dict(sorted(motif_totals.items(), key=lambda item: item[1], reverse=True)),
        "strategy_posture": dict(sorted(strategy_posture.items(), key=lambda item: item[1]["cohort"], reverse=True)),
        "zone_summary": dict(sorted(zone_summary.items(), key=lambda item: item[1]["cohort"], reverse=True)),
        "trust_summary": dict(sorted(trust_summary.items(), key=lambda item: item[1]["signals"], reverse=True)),
        "trust_readout": dict(sorted(trust_readout.items(), key=lambda item: item[1]["signals"], reverse=True)),
        "skill_read": {
            "total_rows": skill_total,
            "rows_with_parent_join": skill_with_parent,
            "parent_join_share": skill_with_parent / skill_total if skill_total else 0,
            "rows": skill_rows,
        },
        "outcome_checks": outcome_checks,
        "tsdr_checks": tsdr_checks,
        "trust_gap": trust_gap,
    }


def color_for_tier(tier: str) -> str:
    return {
        "HIGH_VOLUME_ASSISTIVE_SURFACE": "#2463EB",
        "POST_FRICTION_CONTINUATION": "#0F8A45",
        "EXECUTION_LINKED_WORKFLOW": "#0E7490",
        "SEARCH_TO_AGENT_ESCALATION": "#7C3AED",
        "WEAK_LINKAGE_CONTEXT": "#C2410C",
        "VERIFICATION_ATTACHED_WORKFLOW": "#15803D",
        "OTHER_LINKED_CONTEXT": "#64748B",
        "TRUST_EVIDENCE_GAP": "#C2410C",
        "SCALE_CANDIDATE": "#0F8A45",
        "SHALLOW_ADOPTION": "#D97706",
        "FOCUSED_EXPERT_USE": "#7C3AED",
    }.get(tier, "#475569")


def display_label(label: str) -> str:
    text = label.replace("_", " ").title()
    return text.replace("Trust Attribution ", "")


def load_font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            pass
    return ImageFont.load_default()


def draw_bar_chart(path: Path, title: str, data: dict[str, float], subtitle: str = "") -> None:
    width, height = 1600, 900
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(52, True)
    sub_font = load_font(26)
    label_font = load_font(23, True)
    value_font = load_font(24)
    draw.text((70, 48), title, fill="#061A44", font=title_font)
    if subtitle:
        draw.text((72, 112), subtitle, fill="#334155", font=sub_font)
    items = [(k, v) for k, v in data.items() if v > 0]
    max_value = max([v for _, v in items] or [1])
    top = 185
    bar_h = 58
    gap = 30
    x0 = 470
    x1 = 1420
    for index, (label, value) in enumerate(items[:8]):
        y = top + index * (bar_h + gap)
        label_text = display_label(label)
        draw.text((70, y + 10), label_text, fill="#0F172A", font=label_font)
        draw.rounded_rectangle((x0, y, x1, y + bar_h), radius=18, fill="#E2E8F0")
        bw = int((x1 - x0) * value / max_value)
        draw.rounded_rectangle((x0, y, x0 + bw, y + bar_h), radius=18, fill=color_for_tier(label))
        value_text = fmt_int(value)
        value_box = draw.textbbox((0, 0), value_text, font=value_font)
        value_width = value_box[2] - value_box[0]
        value_x = min(x0 + bw + 18, width - value_width - 70)
        draw.text((value_x, y + 13), value_text, fill="#0F172A", font=value_font)
    img.save(path)


def draw_matrix(path: Path, title: str, rows: list[tuple[str, str, str]]) -> None:
    width, height = 1600, 900
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(52, True)
    head_font = load_font(26, True)
    body_font = load_font(25)
    small_font = load_font(22)
    draw.text((70, 48), title, fill="#061A44", font=title_font)
    y = 150
    for label, status, note in rows:
        color = "#0F8A45" if status == "READY" else "#D97706" if status == "PARTIAL" else "#C2410C"
        draw.rounded_rectangle((70, y, 1530, y + 88), radius=22, fill="#F8FAFC", outline="#CBD5E1", width=2)
        draw.ellipse((100, y + 24, 140, y + 64), fill=color)
        draw.text((165, y + 18), label, fill="#0F172A", font=head_font)
        draw.text((620, y + 18), status, fill=color, font=head_font)
        draw.text((165, y + 52), note, fill="#334155", font=small_font)
        y += 106
    img.save(path)


def write_markdown(summary: dict, out_dir: Path) -> None:
    motif_total = sum(summary["motif_totals"].values())
    zone = summary["zone_summary"]
    trust_gap = summary["trust_gap"]
    lines = [
        "# FluencyTracr Internal Pilot Packet",
        "",
        f"Run date: {summary['generated_on']}",
        "",
        "Status: `INTERNAL_PILOT_REHEARSAL`",
        "",
        "This packet uses company-owned aggregate dogfood exports and a fresh BigQuery run. It is designed to show what FluencyTracr can deliver without asking a client for additional inputs at the start of a pilot.",
        "",
        "## Executive Readout",
        "",
        "FluencyTracr can already produce a meaningful aggregate AI work evidence package from observed telemetry. The strongest current value is not ROI calculation; it is separating activity volume from workflow evidence, trust evidence, source coverage, and value-investigation readiness.",
        "",
        "The internal data shows three distinct realities:",
        "",
        "- AI is broadly present through high-volume assistive surfaces.",
        "- A smaller but more valuable lane shows AI attached to workflow execution, post-friction continuation, and verification evidence.",
        "- The largest commercialization blocker is not lack of activity; it is missing or weak downstream evidence that would connect behavior to business outcomes.",
        "",
        "## Fresh Data Run",
        "",
        f"- BigQuery job: `{summary['fresh_job']['job_id']}`",
        f"- Duration: `{summary['fresh_job']['duration']}`",
        f"- Bytes processed: `{fmt_int(summary['fresh_job']['bytes_processed'])}`",
        f"- Source: {summary['fresh_job']['source']}",
        "",
        "### Microcosm Framework Results",
        "",
        "| Window | Sampled aggregate windows | Frequency | Engagement | Breadth | Reliability | Quality multiplier |",
        "| --- | ---: | ---: | ---: | ---: | ---: | ---: |",
    ]
    for row in summary["framework"]:
        lines.append(
            f"| {row['window']} | {row['sampled_windows']} | {row['avg_frequency']:.1f} | {row['avg_engagement']:.1f} | {row['avg_breadth']:.1f} | {row['avg_reliability']:.2f} | {row['avg_quality']:.2f} |"
        )
    lines.extend(
        [
            "",
            "### Motif Tier Distribution",
            "",
            "| Tier | Aggregate motifs | Share | Interpretation |",
            "| --- | ---: | ---: | --- |",
        ]
    )
    interpretations = {
        "HIGH_VOLUME_ASSISTIVE_SURFACE": "Reach and surface coverage; weak workflow evidence by itself.",
        "POST_FRICTION_CONTINUATION": "Recovery-like evidence: work continues after friction without claiming intent.",
        "EXECUTION_LINKED_WORKFLOW": "AI is attached to actual workflow execution.",
        "SEARCH_TO_AGENT_ESCALATION": "Navigation from search into agent context; not outcome evidence yet.",
        "WEAK_LINKAGE_CONTEXT": "Source-coverage caveat lane.",
        "VERIFICATION_ATTACHED_WORKFLOW": "Strongest trust-adjacent workflow lane.",
        "OTHER_LINKED_CONTEXT": "Residual linked context.",
    }
    for tier, count in summary["motif_totals"].items():
        lines.append(f"| {tier.replace('_', ' ').title()} | {fmt_int(count)} | {fmt_pct(count / motif_total)} | {interpretations.get(tier, '')} |")
    lines.extend(
        [
            "",
            "## Value Realization Readiness",
            "",
            "| Readout zone | Aggregate cohort rows | Signal rows | Attributed signal rows | Executive action |",
            "| --- | ---: | ---: | ---: | --- |",
        ]
    )
    action = {
        "TRUST_EVIDENCE_GAP": "Repair proof loops before value interpretation.",
        "SCALE_CANDIDATE": "Scale-and-measure review; attach outcome evidence.",
        "SHALLOW_ADOPTION": "Workflow redesign before scale.",
        "FOCUSED_EXPERT_USE": "Package the expert pattern if the business confirms value.",
        "SUPPRESSED": "Do not interpret.",
    }
    for key, value in zone.items():
        lines.append(f"| {key.replace('_', ' ').title()} | {fmt_int(value['cohort'])} | {fmt_int(value['signals'])} | {fmt_int(value['attributed'])} | {action.get(key, '')} |")
    lines.extend(
        [
            "",
            "## Trust And Source Coverage",
            "",
            f"- Seven-business-day trust pilot: {fmt_int(trust_gap['high_confidence_episodes'])} high-confidence aggregate product episodes.",
            f"- High-confidence coverage: {fmt_pct(trust_gap['high_confidence_episode_share'])}.",
            f"- Recovered-after-failure pattern inside high-confidence coverage: {fmt_pct(trust_gap['recovered_after_failure_share'])}.",
            f"- Public evidence gap: {fmt_int(trust_gap['public_gap_episodes'])} aggregate episodes ({fmt_pct(trust_gap['public_gap_share'])}).",
            f"- True downstream-evidence gap: {fmt_int(trust_gap['true_downstream_gap'])} aggregate episodes.",
            f"- Ambiguous-boundary fold-in: {fmt_int(trust_gap['ambiguous_boundary_fold_in'])} aggregate episodes.",
            "",
            "Trust evidence is strongest when verification or continuation attaches to workflow-linked paths. Citation clicks and explicit feedback are useful corroboration, but they are too sparse or too attribution-sensitive to carry the trust story alone.",
            "",
            "## Outcome And ROI Boundary",
            "",
            "The current company-owned telemetry is strong enough for value-investigation routing, but not for ROI calculation.",
            "",
            "- Outcome metric identity: missing in saved V4 exports.",
            "- Window-aligned outcome export: held.",
            "- Customer-owned assumptions: missing.",
            "- Behavior-to-outcome attribution: held.",
            "- Time-saved range values: suppressed.",
            "",
            "This should be presented as an evidence-readiness product, not a dollarized ROI product.",
            "",
            "## Pilot Decision Memo",
            "",
            "Decision: continue to a customer-facing rehearsal packet, but keep the claim narrow.",
            "",
            "Recommended posture:",
            "",
            "1. Lead with aggregate AI work evidence, not adoption dashboards.",
            "2. Show motif tiers so executives can distinguish reach, workflow execution, recovery-like continuation, verification, and weak-linkage gaps.",
            "3. Use Velocity x Depth as the operating map.",
            "4. Use Trust Evidence Gap as the proof-loop repair agenda.",
            "5. Treat ROI as blocked until outcome metrics and assumptions are available inside the governed evidence layer.",
            "",
            "Non-goals: no individual scoring, no team ranking, no manager ranking, no productivity claim, no causality claim, no ROI calculation, no raw prompt or output inspection.",
            "",
            "## Glossary",
            "",
        ]
    )
    for term, definition in GLOSSARY:
        lines.append(f"- **{term}:** {definition}")
    (out_dir / "FluencyTracr_Internal_Pilot_Report.md").write_text("\n".join(lines) + "\n")


def write_decision_memo(summary: dict, out_dir: Path) -> None:
    zone = summary["zone_summary"]
    lines = [
        "# Pilot Decision Memo",
        "",
        f"Date: {summary['generated_on']}",
        "",
        "Decision: `PROCEED_TO_EXECUTIVE_REHEARSAL_PACKET`",
        "",
        "## Why",
        "",
        "The internal pilot shows that FluencyTracr can separate raw AI activity from workflow-integrated evidence, trust evidence, source coverage, and value-readiness using aggregate telemetry we already collect.",
        "",
        "## Evidence",
        "",
        f"- Scale-candidate cohort rows: {fmt_int(zone.get('SCALE_CANDIDATE', {}).get('cohort', 0))}.",
        f"- Trust-evidence-gap cohort rows: {fmt_int(zone.get('TRUST_EVIDENCE_GAP', {}).get('cohort', 0))}.",
        f"- Shallow-adoption cohort rows: {fmt_int(zone.get('SHALLOW_ADOPTION', {}).get('cohort', 0))}.",
        f"- Focused-expert-use cohort rows: {fmt_int(zone.get('FOCUSED_EXPERT_USE', {}).get('cohort', 0))}.",
        f"- Skill-read evidence rows with parent joins: {fmt_int(summary['skill_read']['rows_with_parent_join'])} ({fmt_pct(summary['skill_read']['parent_join_share'])}).",
        "",
        "## What We Can Promise",
        "",
        "- Identify where AI is present versus embedded in workflow execution.",
        "- Identify where post-friction continuation and verification evidence exist.",
        "- Identify where trust and source coverage are too weak for interpretation.",
        "- Route value hypotheses to the right next investigation.",
        "",
        "## What We Cannot Promise Yet",
        "",
        "- Dollarized ROI.",
        "- Causal impact.",
        "- Productivity lift.",
        "- Output correctness.",
        "- Employee, team, manager, or department scoring.",
        "",
        "## Next Action",
        "",
        "Build the external-facing executive rehearsal deck/report from this packet, then decide whether to pilot with a customer using the same aggregate-only variable names and evidence lanes.",
        "",
        "## Glossary",
        "",
    ]
    for term, definition in GLOSSARY:
        lines.append(f"- **{term}:** {definition}")
    (out_dir / "FluencyTracr_Pilot_Decision_Memo.md").write_text("\n".join(lines) + "\n")


def add_doc_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(6, 26, 68)


def add_doc_paragraph(document: Document, text: str, bold_lead: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    if bold_lead:
        run.bold = True


def write_docx_from_markdown(md_path: Path, docx_path: Path, title: str) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_para.add_run(title)
    title_run.font.name = "Arial"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(6, 26, 68)
    document.add_paragraph("Internal pilot rehearsal. Aggregate evidence only.")
    for line in md_path.read_text().splitlines()[1:]:
        if not line.strip():
            continue
        if line.startswith("## "):
            add_doc_heading(document, line[3:], 1)
        elif line.startswith("### "):
            add_doc_heading(document, line[4:], 2)
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style=None)
            paragraph.style = styles["Normal"]
            paragraph.paragraph_format.left_indent = Inches(0.18)
            run = paragraph.add_run("• " + line[2:])
            run.font.name = "Arial"
            run.font.size = Pt(10)
        elif line.startswith("| "):
            # Keep markdown tables as compact monospaced text for this internal
            # packet; the deck carries the visual charts.
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(7.8)
        elif line.startswith("# "):
            continue
        else:
            add_doc_paragraph(document, line)
    document.save(docx_path)


def build_packet(input_dir: Path, out_dir: Path) -> None:
    ensure_dir(out_dir)
    chart_dir = out_dir / "charts"
    ensure_dir(chart_dir)
    summary = load_summary(input_dir)
    summary_path = out_dir / "pilot_packet_summary.json"
    summary_path.write_text(json.dumps(summary, indent=2))

    draw_bar_chart(
        chart_dir / "motif_tier_distribution.png",
        "Aggregate Motif Tiers",
        summary["motif_totals"],
        "Fresh 90-day internal microcosm run; aggregate motifs only.",
    )
    draw_bar_chart(
        chart_dir / "readout_zone_distribution.png",
        "Value-Readiness Zones",
        {key: value["cohort"] for key, value in summary["zone_summary"].items()},
        "Existing V4 aggregate dogfood exports.",
    )
    draw_bar_chart(
        chart_dir / "trust_attribution_distribution.png",
        "Trust Signal Attribution",
        {key: value["signals"] for key, value in summary["trust_summary"].items()},
        "Attribution status across retained trust-signal summary.",
    )
    draw_matrix(
        chart_dir / "data_readiness_matrix.png",
        "Data Readiness",
        [
            ("AI surfaces and workflow events", "READY", "Observed in scrubbed GCE events and agent spans."),
            ("Motif tiers and Velocity x Depth", "READY", "Fresh run plus retained V4 aggregate exports."),
            ("Trust verification and feedback", "PARTIAL", "Available, but attribution gaps remain material."),
            ("Skill-read evidence", "READY", f"{fmt_pct(summary['skill_read']['parent_join_share'])} of retained skill-read rows have parent join keys."),
            ("Outcome and assumption data", "MISSING", "No governed outcome metric or assumption ledger in saved exports."),
            ("ROI calculation", "MISSING", "Blocked by design until outcome evidence and assumptions exist."),
        ],
    )
    write_markdown(summary, out_dir)
    write_decision_memo(summary, out_dir)
    write_docx_from_markdown(
        out_dir / "FluencyTracr_Internal_Pilot_Report.md",
        out_dir / "FluencyTracr_Internal_Pilot_Report.docx",
        "FluencyTracr Internal Pilot Report",
    )
    write_docx_from_markdown(
        out_dir / "FluencyTracr_Pilot_Decision_Memo.md",
        out_dir / "FluencyTracr_Pilot_Decision_Memo.docx",
        "FluencyTracr Pilot Decision Memo",
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-dir", default="dogfood-output/internal-pilot-packet-2026-05-28")
    parser.add_argument("--output-dir", default="output/internal-pilot-packet-2026-05-28")
    args = parser.parse_args()
    build_packet(ROOT / args.input_dir, ROOT / args.output_dir)


if __name__ == "__main__":
    main()
